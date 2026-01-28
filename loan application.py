import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document 
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageTk  # Required for previews
import database 
import datetime
import uuid
import subprocess
import sys
import os

# --- THEME & STYLE ---
PRIMARY_GREEN = "#2ecc71"
PRIMARY_BLUE = "#2980b9"
BG_LIGHT = "#f4f7f6"
DARK_TEXT = "#2c3e50"
BORDER_COLOR = "#dcdde1"
DANGER_RED = "#c0392b"
HOVER_RED = "#e74c3c"
FONT_FAMILY = "Segoe UI" 

# --- SESSION PERSISTENCE ---
try:
    CURRENT_USER_ROLE = sys.argv[1]
    CURRENT_USER_NAME = sys.argv[2]
except IndexError:
    CURRENT_USER_ROLE = "Staff"
    CURRENT_USER_NAME = "Guest"

class LoanApplicationApp:
    def __init__(self, root):
        self.root = root
        self.root.title(f"Loan Application - Logged in as: {CURRENT_USER_NAME}")
        self.root.geometry("1150x850") 
        self.root.configure(bg=BG_LIGHT)

        try:
            title_icon = tk.PhotoImage(file="bu logo.png")
            self.root.iconphoto(True, title_icon)
        except Exception:
            pass

        self.repayment_method_var = tk.StringVar(value="Monthly")
        self.terms_var = tk.IntVar()
        self.security_photos = [] # Store file paths
        self.preview_images = [] # Keep reference to images to prevent garbage collection
        
        self.setup_ui()

    def setup_ui(self):
        header = tk.Frame(self.root, bg=PRIMARY_GREEN, height=100)
        header.pack(fill="x", side="top")
        header.pack_propagate(False)

        tk.Label(header, text="OFFICIAL LOAN APPLICATION", font=(FONT_FAMILY, 24, "bold"), 
                 bg=PRIMARY_GREEN, fg="white").pack(pady=(25, 0))

        self.main_canvas = tk.Canvas(self.root, bg=BG_LIGHT, highlightthickness=0)
        self.scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=self.main_canvas.yview)
        self.scrollable_frame = tk.Frame(self.main_canvas, bg=BG_LIGHT)

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.main_canvas.configure(scrollregion=self.main_canvas.bbox("all"))
        )

        self.main_canvas.create_window((575, 0), window=self.scrollable_frame, anchor="n")
        self.main_canvas.configure(yscrollcommand=self.scrollbar.set)
        self.main_canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")

        self.card = tk.Frame(self.scrollable_frame, bg="white", padx=50, pady=40, 
                             highlightthickness=1, highlightbackground=BORDER_COLOR)
        self.card.pack(pady=30, padx=20)

        self.card.grid_columnconfigure(0, weight=1)
        self.card.grid_columnconfigure(1, weight=1)

        self.build_form()

    def create_label(self, text, row, col, colspan=1):
        lbl = tk.Label(self.card, text=text, font=(FONT_FAMILY, 11, "bold"), bg="white", fg=DARK_TEXT)
        lbl.grid(row=row, column=col, columnspan=colspan, sticky="w", padx=15, pady=(15, 5))

    def create_entry(self, row, col, colspan=1):
        ent = tk.Entry(self.card, font=(FONT_FAMILY, 14), bd=0, highlightthickness=1, 
                       highlightbackground=BORDER_COLOR)
        ent.grid(row=row, column=col, columnspan=colspan, sticky="ew", padx=15, pady=(0, 20), ipady=12)
        return ent

    def attach_photos(self):
        files = filedialog.askopenfilenames(
            title="Select Security Photos",
            filetypes=[("Image files", "*.jpg *.jpeg *.png *.bmp")]
        )
        if files:
            new_files = list(files)
            if (len(self.security_photos) + len(new_files)) > 5:
                messagebox.showwarning("Limit Exceeded", "You can only attach a maximum of 5 photos in total.")
                self.security_photos.extend(new_files[:5 - len(self.security_photos)])
            else:
                self.security_photos.extend(new_files)
            
            self.update_photo_label()

    def update_photo_label(self):
        count = len(self.security_photos)
        if count > 0:
            self.photo_count_lbl.config(text=f"📎 {count} Photos Attached", fg=PRIMARY_GREEN)
        else:
            self.photo_count_lbl.config(text="No photos attached", fg="#7f8c8d")

    def remove_photo(self, path, frame, win):
        if path in self.security_photos:
            self.security_photos.remove(path)
            frame.destroy()
            self.update_photo_label()
            if not self.security_photos:
                win.destroy()

    def preview_photos(self):
        if not self.security_photos:
            messagebox.showinfo("Preview", "No photos attached to preview.")
            return

        preview_win = tk.Toplevel(self.root)
        preview_win.title("Security Photo Previews")
        preview_win.geometry("900x500")
        preview_win.configure(bg=BG_LIGHT)

        self.preview_images = [] # Clear old references
        
        container = tk.Frame(preview_win, bg=BG_LIGHT, padx=20, pady=20)
        container.pack(fill="both", expand=True)

        for path in self.security_photos:
            try:
                img = Image.open(path)
                img.thumbnail((200, 180)) 
                tk_img = ImageTk.PhotoImage(img)
                self.preview_images.append(tk_img) 

                photo_frame = tk.Frame(container, bg="white", bd=1, relief="solid", padx=5, pady=5)
                photo_frame.pack(side="left", padx=10, pady=10, anchor="n")
                
                lbl = tk.Label(photo_frame, image=tk_img, bg="white")
                lbl.pack()
                
                # Truncate long filenames
                fname = os.path.basename(path)
                short_name = (fname[:15] + '..') if len(fname) > 17 else fname
                tk.Label(photo_frame, text=short_name, font=(FONT_FAMILY, 8), bg="white").pack(pady=2)

                del_btn = tk.Button(photo_frame, text="❌ REMOVE", bg=DANGER_RED, fg="white", 
                                    font=(FONT_FAMILY, 7, "bold"), bd=0, cursor="hand2",
                                    command=lambda p=path, f=photo_frame: self.remove_photo(p, f, preview_win))
                del_btn.pack(fill="x", pady=2)

            except Exception as e:
                tk.Label(container, text=f"Error: {os.path.basename(path)}", fg="red").pack()

    def build_form(self):
        self.create_label("NIN NUMBER (NATIONAL ID)", 0, 0)
        
        nin_container = tk.Frame(self.card, bg="white")
        nin_container.grid(row=1, column=0, sticky="ew", padx=15, pady=(0, 20))
        
        self.nin_entry = tk.Entry(nin_container, font=(FONT_FAMILY, 14), bd=0, highlightthickness=1, 
                                 highlightbackground=BORDER_COLOR)
        self.nin_entry.pack(side="left", fill="x", expand=True, ipady=12)
        
        search_btn = tk.Button(nin_container, text="🔍 FETCH", bg=PRIMARY_BLUE, fg="white", 
                               font=(FONT_FAMILY, 9, "bold"), command=self.lookup_customer, cursor="hand2")
        search_btn.pack(side="right", padx=(5, 0), ipady=8)

        self.create_label("FULL NAME OF APPLICANT", 0, 1)
        self.name_entry = self.create_entry(1, 1)

        self.create_label("LOAN AMOUNT (RWF)", 2, 0)
        self.create_label("LOAN CATEGORY", 2, 1)
        self.amount_entry = self.create_entry(3, 0)
        self.amount_entry.bind("<KeyRelease>", self.update_return_amount)
        
        self.type_combo = ttk.Combobox(self.card, values=["Personal", "Business", "Home", "Education", "Vehicle"], 
                                        font=(FONT_FAMILY, 13), state="readonly")
        self.type_combo.grid(row=3, column=1, sticky="ew", padx=15, pady=(0, 20))

        self.create_label("REPAYMENT DURATION", 4, 0)
        self.create_label("COLLATERAL SECURITY", 4, 1)
        self.duration_combo = ttk.Combobox(self.card, values=["6 months", "1 year", "2 years", "3 years", "5 years"], 
                                            font=(FONT_FAMILY, 13), state="readonly")
        self.duration_combo.grid(row=5, column=0, sticky="ew", padx=15, pady=(0, 20))
        self.duration_combo.bind("<<ComboboxSelected>>", self.update_return_amount)

        collateral_frame = tk.Frame(self.card, bg="white")
        collateral_frame.grid(row=5, column=1, sticky="ew", padx=15, pady=(0, 20))
        
        self.collateral_combo = ttk.Combobox(collateral_frame, values=["Land Title", "Vehicle Logbook", "House Property", "Equipment", "Guarantor", "Machinery and equipment", "Salary assignment"], 
                                            font=(FONT_FAMILY, 13), state="readonly")
        self.collateral_combo.pack(side="top", fill="x")
        
        photo_actions = tk.Frame(collateral_frame, bg="white")
        photo_actions.pack(side="top", fill="x", pady=(5,0))

        self.photo_btn = tk.Button(photo_actions, text="📸 ATTACH", bg="#95a5a6", fg="white",
                                   font=(FONT_FAMILY, 9, "bold"), command=self.attach_photos, cursor="hand2", width=12)
        self.photo_btn.pack(side="left", padx=(0, 5))

        self.preview_btn = tk.Button(photo_actions, text="👁 PREVIEW", bg=PRIMARY_BLUE, fg="white",
                                     font=(FONT_FAMILY, 9, "bold"), command=self.preview_photos, cursor="hand2", width=12)
        self.preview_btn.pack(side="left")
        
        self.photo_count_lbl = tk.Label(collateral_frame, text="No photos attached", font=(FONT_FAMILY, 8), bg="white", fg="#7f8c8d")
        self.photo_count_lbl.pack(side="top", anchor="w")

        self.create_label("PAYMENT FREQUENCY", 6, 0)
        radio_frame = tk.Frame(self.card, bg="white")
        radio_frame.grid(row=7, column=0, sticky="w", padx=15)
        tk.Radiobutton(radio_frame, text="Monthly", variable=self.repayment_method_var, value="Monthly", bg="white", font=(FONT_FAMILY, 12)).pack(side="left")
        tk.Radiobutton(radio_frame, text="Weekly", variable=self.repayment_method_var, value="Weekly", bg="white", font=(FONT_FAMILY, 12)).pack(side="left", padx=20)

        self.create_label("PURPOSE OF LOAN", 8, 0, colspan=2)
        self.purpose_text = tk.Text(self.card, height=4, font=(FONT_FAMILY, 12), bd=1, relief="solid", padx=10, pady=10)
        self.purpose_text.grid(row=9, column=0, columnspan=2, sticky="ew", padx=15, pady=(0, 25))

        self.total_frame = tk.Frame(self.card, bg="#f1f2f6", padx=30, pady=25)
        self.total_frame.grid(row=10, column=0, columnspan=2, sticky="ew", padx=15, pady=10)
        tk.Label(self.total_frame, text="ESTIMATED TOTAL REPAYMENT (12% Interest)", font=(FONT_FAMILY, 10, "bold"), bg="#f1f2f6", fg=DARK_TEXT).pack(anchor="w")
        self.return_amount_lbl = tk.Label(self.total_frame, text="0.00 RWF", font=(FONT_FAMILY, 28, "bold"), bg="#f1f2f6", fg=PRIMARY_GREEN)
        self.return_amount_lbl.pack(anchor="w")

        tk.Checkbutton(self.card, text="I accept the terms and conditions", variable=self.terms_var, bg="white", font=(FONT_FAMILY, 11)).grid(row=11, column=0, columnspan=2, sticky="w", padx=15, pady=20)

        submit_btn_frame = tk.Frame(self.card, bg="white")
        submit_btn_frame.grid(row=12, column=0, columnspan=2, pady=(10, 20))

        tk.Button(submit_btn_frame, text="SUBMIT TO DATABASE", bg=PRIMARY_GREEN, fg="white", font=(FONT_FAMILY, 11, "bold"), bd=0, width=20, height=2, cursor="hand2", command=self.submit_application).pack(side="left", padx=10)
        tk.Button(submit_btn_frame, text="PRINT WORD DOC", bg="#3498db", fg="white", font=(FONT_FAMILY, 11, "bold"), bd=0, width=18, height=2, cursor="hand2", command=self.print_application).pack(side="left", padx=10)

        footer_nav = tk.Frame(self.scrollable_frame, bg=BG_LIGHT)
        footer_nav.pack(fill="x", pady=(20, 50))

        tk.Button(footer_nav, text="🔙 BACK TO DASHBOARD", bg=PRIMARY_BLUE, fg="white", font=(FONT_FAMILY, 12, "bold"), 
                  bd=0, width=25, height=2, cursor="hand2", command=self.return_to_dashboard).pack(side="left", padx=(150, 20))

        self.logout_btn = tk.Button(footer_nav, text="🛑 LOGOUT SYSTEM", bg=DANGER_RED, fg="white", font=(FONT_FAMILY, 12, "bold"), 
                                    bd=0, width=25, height=2, cursor="hand2", command=self.handle_logout)
        self.logout_btn.pack(side="left", padx=20)

    def lookup_customer(self):
        nin = self.nin_entry.get().strip()
        if not nin:
            messagebox.showwarning("Input Required", "Please enter a NIN number to search.")
            return
        try:
            prev_record = database.db['loans'].find_one({"nin_number": nin}, sort=[("application_date", -1)])
            if prev_record:
                self.name_entry.delete(0, tk.END)
                self.name_entry.insert(0, prev_record.get("customer_name", ""))
                collateral_val = prev_record.get("collateral", "")
                if collateral_val in self.collateral_combo['values']:
                    self.collateral_combo.set(collateral_val)
                messagebox.showinfo("User Found", f"Records found for {prev_record.get('customer_name')}.")
            else:
                messagebox.showinfo("New Customer", "No existing records found for this NIN.")
        except Exception as e:
            messagebox.showerror("Search Error", f"Could not retrieve records: {e}")

    def return_to_dashboard(self):
        self.root.destroy()
        try:
            subprocess.Popen([sys.executable, "dashboard.py", CURRENT_USER_ROLE, CURRENT_USER_NAME])
        except Exception:
            messagebox.showerror("Error", "Could not find 'dashboard.py'.")

    def handle_logout(self):
        if messagebox.askyesno("Logout", "Are you sure you want to sign out?"):
            database.log_activity(CURRENT_USER_NAME, "Logout", "User signed out from Application form")
            self.root.destroy()
            try:
                subprocess.Popen([sys.executable, "login.py"])
            except Exception:
                messagebox.showerror("Error", "Could not find 'login.py'.")

    def update_return_amount(self, event=None):
        try:
            amt_str = self.amount_entry.get().replace(',', '')
            amt = float(amt_str) if amt_str else 0
            dur = self.duration_combo.get()
            years = int(dur.split()[0]) if "year" in dur else int(dur.split()[0]) / 12.0
            total = amt + (amt * 0.12 * years)
            self.return_amount_lbl.config(text=f"{total:,.2f} RWF")
            return total
        except:
            self.return_amount_lbl.config(text="0.00 RWF")
            return 0

    def submit_application(self):
        current_name = self.name_entry.get().strip()
        current_nin = self.nin_entry.get().strip()
        
        if not current_name or not self.amount_entry.get() or not current_nin:
            messagebox.showerror("Error", "Please fill in all required fields.")
            return
        if self.terms_var.get() == 0:
            messagebox.showwarning("Terms", "Please accept the terms.")
            return

        try:
            existing_user = database.db['loans'].find_one({"nin_number": current_nin})
            
            if existing_user:
                stored_name = existing_user.get("customer_name", "").strip().lower()
                if stored_name != current_name.lower():
                    messagebox.showerror("Identity Error", 
                        f"Registration Failed!\n\nThis NIN ({current_nin}) is already associated with a different name: "
                        f"'{existing_user.get('customer_name')}'.\n\nOne NIN cannot be used by more than one person.")
                    return

            current_year = datetime.datetime.now().strftime("%Y")
            unique_suffix = str(uuid.uuid4())[:4].upper()
            loan_id = f"LOAN-{current_year}-{unique_suffix}"
            
            loan_data = {
                "loan_id": loan_id,
                "customer_name": current_name,
                "nin_number": current_nin,
                "loan_amount": float(self.amount_entry.get().replace(',', '')),
                "loan_type": self.type_combo.get(),
                "duration": self.duration_combo.get(),
                "collateral": self.collateral_combo.get(),
                "security_photos": self.security_photos,
                "payment_plan": self.repayment_method_var.get(),
                "purpose": self.purpose_text.get("1.0", tk.END).strip(),
                "return_amount": self.update_return_amount(),
                "status": "Pending",
                "application_date": datetime.datetime.now()
            }
            database.db['loans'].insert_one(loan_data)
            database.log_activity(CURRENT_USER_NAME, "New Loan Application", f"Submitted loan {loan_id} for {current_name}")
            
            messagebox.showinfo("Success", f"Application {loan_id} saved to Database!")
            if messagebox.askyesno("Print", "Generate Word Doc for signing?"):
                self.print_application(custom_id=loan_id)
            self.return_to_dashboard()
        except Exception as e:
            messagebox.showerror("System Error", f"Failed to save: {e}")

    def print_application(self, custom_id=None):
        if not self.name_entry.get().strip() or not self.amount_entry.get().strip():
            messagebox.showwarning("Incomplete Form", "Missing info for printing.")
            return

        try:
            app_id = custom_id if custom_id else "TEMP-" + str(uuid.uuid4())[:5]
            file_path = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=f"Loan_App_{app_id}.docx")
            if not file_path: return

            doc = Document()
            try:
                doc.add_picture('bu logo.png', width=Inches(1.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except: pass

            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("OFFICIAL LOAN APPLICATION FORM")
            run.bold = True
            run.font.size = Pt(20)

            ref_p = doc.add_paragraph()
            ref_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ref_p.add_run(f"Reference ID: {app_id} | Date: {datetime.datetime.now().strftime('%Y-%m-%d')}")
            doc.add_paragraph("_" * 75)

            doc.add_heading('I. APPLICANT INFORMATION', level=2)
            p1 = doc.add_paragraph()
            p1.add_run("Full Name: ").bold = True; p1.add_run(self.name_entry.get().upper())
            p2 = doc.add_paragraph()
            p2.add_run("National ID (NIN): ").bold = True; p2.add_run(self.nin_entry.get())

            doc.add_heading('II. LOAN SPECIFICATIONS', level=2)
            table = doc.add_table(rows=0, cols=2)
            specs = [
                ("Requested Amount:", f"{self.amount_entry.get()} RWF"),
                ("Loan Category:", self.type_combo.get()),
                ("Repayment Term:", self.duration_combo.get()),
                ("Collateral Provided:", self.collateral_combo.get()),
                ("Repayment Cycle:", self.repayment_method_var.get()),
                ("Interest Rate:", "12% Per Annum")
            ]

            for label, val in specs:
                row_cells = table.add_row().cells
                row_cells[0].text = label; row_cells[0].paragraphs[0].runs[0].bold = True
                row_cells[1].text = val

            doc.add_paragraph()
            summary = doc.add_paragraph()
            res_run = summary.add_run(f"TOTAL ESTIMATED REPAYMENT: {self.return_amount_lbl.cget('text')}")
            res_run.bold = True; res_run.font.size = Pt(13)

            doc.add_heading('III. PURPOSE OF LOAN', level=2)
            doc.add_paragraph(self.purpose_text.get("1.0", tk.END).strip())

            if self.security_photos:
                doc.add_heading('IV. SECURITY ATTACHMENTS', level=2)
                for photo_path in self.security_photos:
                    if os.path.exists(photo_path):
                        try:
                            doc.add_picture(photo_path, width=Inches(3.5))
                            p = doc.add_paragraph(f"Attachment: {os.path.basename(photo_path)}")
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except:
                            doc.add_paragraph(f"[Error loading image: {os.path.basename(photo_path)}]")

            doc.add_paragraph("\n" * 2)
            sig_table = doc.add_table(rows=1, cols=2)
            sig_table.width = doc.sections[0].page_width
            sig_table.cell(0,0).paragraphs[0].add_run("__________________________\nAPPLICANT SIGNATURE")
            right_cell = sig_table.cell(0,1).paragraphs[0]
            right_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_cell.add_run("__________________________\nOFFICER APPROVAL / DATE")

            doc.save(file_path)
            database.log_activity(CURRENT_USER_NAME, "Print Application", f"Generated Word document for {self.name_entry.get()}")
            os.startfile(file_path)
        except Exception as e:
            messagebox.showerror("Print Error", f"Failed to generate document: {e}")

if __name__ == "__main__":
    root = tk.Tk()
    app = LoanApplicationApp(root)
    root.mainloop()